#!/usr/bin/env python3
"""Convert the compiled Markdown manuscript to a .docx suitable for upload to Google Drive."""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC = Path(__file__).parent / "Return_of_the_Citizen_FULL_MANUSCRIPT.md"
DST = Path(__file__).parent / "Return_of_the_Citizen_FULL_MANUSCRIPT.docx"


def parse_inline(paragraph, text):
    """Parse inline italic/bold formatting from markdown and append runs."""
    # Handle bold (**text**) and italic (*text*).
    # Simple stateful parser; no nested bold-italic.
    tokens = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*)', text)
    for tok in tokens:
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            run = paragraph.add_run(tok[2:-2])
            run.bold = True
        elif tok.startswith("*") and tok.endswith("*"):
            run = paragraph.add_run(tok[1:-1])
            run.italic = True
        else:
            paragraph.add_run(tok)


def build():
    doc = Document()

    # Base style.
    style = doc.styles['Normal']
    style.font.name = 'Georgia'
    style.font.size = Pt(11)

    # Page margins.
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.1)
        section.right_margin = Inches(1.1)

    text = SRC.read_text(encoding='utf-8')
    lines = text.split('\n')

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        # Skip empty lines between blocks; we'll add paragraph breaks naturally.
        if not line:
            i += 1
            continue

        # Horizontal rule: volume/section divider.
        if re.fullmatch(r'-{3,}', line):
            # Add a page break only between major volumes; use blank paragraph otherwise.
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run('').font.size = Pt(11)
            i += 1
            continue

        # Scene break in chapters (centered ***).
        if line.strip() == '***':
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run('* * *')
            run.font.size = Pt(11)
            i += 1
            continue

        # Headings.
        if line.startswith('### '):
            p = doc.add_heading(level=3)
            parse_inline(p, line[4:].strip())
            i += 1
            continue
        if line.startswith('## '):
            p = doc.add_heading(level=2)
            parse_inline(p, line[3:].strip())
            i += 1
            continue
        if line.startswith('# '):
            # Major headings: add a page break before.
            doc.add_page_break()
            p = doc.add_heading(level=1)
            parse_inline(p, line[2:].strip())
            i += 1
            continue

        # Accumulate a paragraph: consecutive non-blank, non-special lines.
        buf = [line]
        j = i + 1
        while j < len(lines):
            nxt = lines[j].rstrip()
            if not nxt:
                break
            if re.fullmatch(r'-{3,}', nxt) or nxt.strip() == '***':
                break
            if nxt.startswith('#'):
                break
            buf.append(nxt)
            j += 1
        paragraph_text = ' '.join(buf)
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.first_line_indent = Inches(0.25)
        parse_inline(p, paragraph_text)
        i = j

    doc.save(DST)
    print(f"Wrote {DST} ({DST.stat().st_size:,} bytes)")


if __name__ == '__main__':
    build()
